from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript.md"
OUTPUT = ROOT / "manuscript.docx"


def add_inline_markdown(paragraph, text: str) -> None:
    """Add a small subset of markdown inline formatting to a Word paragraph."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if line.strip()]
    if len(rows) >= 2 and is_table_separator(lines[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header = rows[0]
        body = rows[1:]

    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for idx, cell_text in enumerate(header):
        cell = table.rows[0].cells[idx]
        p = cell.paragraphs[0]
        add_inline_markdown(p, cell_text)
        for run in p.runs:
            run.bold = True

    for row in body:
        cells = table.add_row().cells
        for idx, cell_text in enumerate(row[: len(header)]):
            add_inline_markdown(cells[idx].paragraphs[0], cell_text)

    document.add_paragraph()


def add_image(document: Document, alt_text: str, image_ref: str) -> None:
    image_path = (ROOT / image_ref).resolve()
    if not image_path.exists():
        p = document.add_paragraph()
        p.add_run(f"[Missing figure: {image_ref}]").italic = True
        return
    if alt_text:
        p = document.add_paragraph()
        run = p.add_run(alt_text)
        run.bold = True
    document.add_picture(str(image_path), width=Inches(6.4))
    document.add_paragraph()


def add_code_block(document: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    p = document.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    document.add_paragraph()


def build_docx() -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for level in range(1, 4):
        styles[f"Heading {level}"].font.name = "Calibri"

    lines = MANUSCRIPT.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            add_markdown_table(document, table_lines)
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(document, image_match.group(1), image_match.group(2))
            idx += 1
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
            idx += 1
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            idx += 1
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
            idx += 1
            continue
        if stripped.startswith("#### "):
            document.add_heading(stripped[5:].strip(), level=3)
            idx += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet_match:
            p = document.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet_match.group(1))
            idx += 1
            continue
        if number_match:
            p = document.add_paragraph(style="List Number")
            add_inline_markdown(p, number_match.group(1))
            idx += 1
            continue

        p = document.add_paragraph()
        add_inline_markdown(p, stripped)
        idx += 1

    if code_lines:
        add_code_block(document, code_lines)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("Figure Captions", level=1)
    captions_path = ROOT / "figure_captions.md"
    if captions_path.exists():
        for caption_line in captions_path.read_text(encoding="utf-8").splitlines():
            if caption_line.startswith("- "):
                p = document.add_paragraph(style="List Bullet")
                add_inline_markdown(p, caption_line[2:])

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_docx()
    print(f"Saved {path}")
